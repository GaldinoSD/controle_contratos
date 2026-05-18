# -*- coding: utf-8 -*-

import os
import io
from pathlib import Path
from datetime import datetime, date, time, timedelta
from zoneinfo import ZoneInfo

from flask import (
    Flask, render_template, request, redirect,
    flash, jsonify, url_for, abort
)
from flask_sqlalchemy import SQLAlchemy
from flask_login import (
    LoginManager, UserMixin, login_user, logout_user,
    login_required, current_user
)
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash


# =====================================================
# TIMEZONE (CORRIGE +3H)
# =====================================================
TZ = ZoneInfo("America/Sao_Paulo")

def agora():
    # datetime com timezone correto (BR)
    return datetime.now(TZ)

def parse_periodo_local(periodo_str: str):
    """
    Recebe 'dd/mm/YYYY até dd/mm/YYYY' (horário BR)
    Retorna (inicio, fim) como datetime TZ-aware
    """
    if not periodo_str or " até " not in periodo_str:
        return (None, None)

    try:
        ini_str, fim_str = periodo_str.split(" até ")
        ini = datetime.strptime(ini_str.strip(), "%d/%m/%Y")
        fim = datetime.strptime(fim_str.strip(), "%d/%m/%Y")

        inicio = ini.replace(hour=0, minute=0, second=0, microsecond=0, tzinfo=TZ)
        fim = fim.replace(hour=23, minute=59, second=59, microsecond=999999, tzinfo=TZ)

        return (inicio, fim)
    except Exception:
        return (None, None)


# =====================================================
# SENHA MASTER DO SISTEMA (ATENÇÃO: texto puro)
# =====================================================
MASTER_PASSWORD = "26828021jJ*"


# =====================================================
# CONFIGURAÇÃO DO APP
# =====================================================
app = Flask(__name__)
app.config["SECRET_KEY"] = "secreto"
app.config["SQLALCHEMY_DATABASE_URI"] = os.getenv(
    "DATABASE_URL",
    "postgresql://jonatas:26828021jJ@localhost/contratos")

# PASTAS EXISTENTES
app.config["UPLOAD_FOLDER"] = "static/uploads"
app.config["FOTOS_COLAB"] = "static/fotos"
app.config["LOGO_FOLDER"] = "static/logos"
app.config["MANUAIS_FOLDER"] = "static/manuais"
app.config["TREINAMENTOS_FOLDER"] = "static/uploads/treinamentos"

db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = "login"


# =====================================================
# HELPERS
# =====================================================
def admin_only():
    """Bloqueia acesso se não for admin."""
    if not current_user.is_authenticated:
        abort(401)
    if current_user.role != "admin":
        abort(403)

def ensure_upload_folder():
    """Garante que a pasta de upload exista (evita erro ao salvar arquivo)."""
    os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

def ensure_base_folders():
    """Garante que as pastas do sistema existam."""
    os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
    os.makedirs(app.config["FOTOS_COLAB"], exist_ok=True)
    os.makedirs(app.config["LOGO_FOLDER"], exist_ok=True)
    os.makedirs(app.config["MANUAIS_FOLDER"], exist_ok=True)
    os.makedirs(app.config["TREINAMENTOS_FOLDER"], exist_ok=True)

    # pastas opcionais
    os.makedirs(os.path.join(app.root_path, "static", "inbox"), exist_ok=True)
    os.makedirs(os.path.join(app.root_path, "static", "relatorios"), exist_ok=True)


# ✅ Cria as pastas ao iniciar (recomendado)
ensure_base_folders()



# =====================================================
# MODELOS DO BANCO
# =====================================================

# ==============================================
# USUÁRIOS
# ==============================================
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(120))
    email = db.Column(db.String(120), unique=True)
    password = db.Column(db.String(255))
    role = db.Column(db.String(20))  # admin / colaborador / cliente_colaborador
    company_id = db.Column(db.Integer, db.ForeignKey("company.id"), nullable=True)

    company = db.relationship("Company")

    def set_password(self, password):
        self.password = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password, password)

# ==============================================
# EMPRESAS
# ==============================================
class Company(db.Model):
    __tablename__ = "company"

    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(120), nullable=False)
    cnpj = db.Column(db.String(20), nullable=False)
    created_at = db.Column(db.DateTime, default=agora)
    active = db.Column(db.Boolean, default=True)
    logo = db.Column(db.String(255), nullable=True)

# ==============================================
# CONTRATOS
# ==============================================
class Contract(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    company_id = db.Column(db.Integer, db.ForeignKey("company.id"))
    description = db.Column(db.Text)
    start_date = db.Column(db.String(20))
    end_date = db.Column(db.String(20))
    status = db.Column(db.String(20), default="ativo")

    company = db.relationship("Company")

# ==============================================
# ARQUIVOS DO CONTRATO
# ==============================================
class ContractFile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    contract_id = db.Column(db.Integer, db.ForeignKey("contract.id"))
    file_path = db.Column(db.String(255))
    uploaded_at = db.Column(db.DateTime, default=agora)

# ==============================================
# TAREFAS
# ==============================================
class Task(db.Model):
    __tablename__ = "task"

    id = db.Column(db.Integer, primary_key=True)
    contract_id = db.Column(db.Integer, db.ForeignKey("contract.id"))

    title = db.Column(db.String(120), nullable=False)
    description = db.Column(db.Text)

    # ⚠️ se puder, o ideal é Date, mas mantive String para não quebrar nada
    due_date = db.Column(db.String(20))

    priority = db.Column(db.String(20), default="Normal")

    # STATUS: pendente | andamento | concluida
    status = db.Column(db.String(20), default="pendente", nullable=False)

    # Responsável (User.id) ou None (Todos)
    assigned_to = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=True)

    created_at = db.Column(db.DateTime, default=agora)
    completed_at = db.Column(db.DateTime)

    # =========================
    # RELACIONAMENTOS
    # =========================
    contract = db.relationship(
        "Contract",
        backref=db.backref("tasks", lazy="dynamic")
    )

    assigned_user = db.relationship(
        "User",
        backref=db.backref("tarefas_recebidas", lazy="dynamic")
    )

    # 🔥 ETAPAS
    steps = db.relationship(
        "TaskStep",
        back_populates="task",
        cascade="all, delete-orphan",
        order_by="TaskStep.created_at"
    )

    # Conclusão final
    completion = db.relationship(
        "TaskCompletion",
        backref="task",
        uselist=False,
        cascade="all, delete-orphan"
    )

    # Logs antigos
    logs = db.relationship(
        "TaskLog",
        backref="task",
        cascade="all, delete-orphan"
    )

    # =========================
    # MÉTODOS
    # =========================
    def iniciar(self):
        if self.status == "pendente":
            self.status = "andamento"

    def concluir(self):
        self.status = "concluida"
        self.completed_at = agora()

    def __repr__(self):
        return f"<Task {self.id} - {self.title} ({self.status})>"

# ==============================================
# ETAPAS DA TAREFA
# ==============================================
class TaskStep(db.Model):
    __tablename__ = "task_steps"

    id = db.Column(db.Integer, primary_key=True)

    task_id = db.Column(
        db.Integer,
        db.ForeignKey("task.id"),
        nullable=False
    )

    user_id = db.Column(
        db.Integer,
        db.ForeignKey("user.id"),
        nullable=False
    )

    description = db.Column(db.Text, nullable=False)
    file_path = db.Column(db.String(255))

    created_at = db.Column(
        db.DateTime,
        default=agora,
        nullable=False
    )

    task = db.relationship(
        "Task",
        back_populates="steps"
    )

    user = db.relationship(
        "User",
        backref=db.backref("task_steps", lazy="dynamic")
    )

    def __repr__(self):
        return f"<TaskStep {self.id} | Task {self.task_id} | User {self.user_id}>"

# ==============================================
# CONCLUSÃO DA TAREFA
# ==============================================
class TaskCompletion(db.Model):
    __tablename__ = "task_completion"

    id = db.Column(db.Integer, primary_key=True)

    task_id = db.Column(
        db.Integer,
        db.ForeignKey("task.id"),
        nullable=False,
        unique=True
    )

    user = db.Column(db.String(120))  # nome do colaborador
    note = db.Column(db.Text, nullable=False)
    file_path = db.Column(db.String(255))
    created_at = db.Column(db.DateTime, default=agora)

    def __repr__(self):
        return f"<TaskCompletion Task {self.task_id}>"

# ==============================================
# LOG ANTIGO (mantido por compatibilidade)
# ==============================================
class TaskLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    task_id = db.Column(db.Integer, db.ForeignKey("task.id"))
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"))
    note = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=agora)

    user = db.relationship("User")
    files = db.relationship("TaskFile", backref="log", cascade="all, delete-orphan")

class TaskFile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    task_log_id = db.Column(db.Integer, db.ForeignKey("task_log.id"))
    file_path = db.Column(db.String(255))
    uploaded_at = db.Column(db.DateTime, default=agora)

# ==============================================
# COLABORADORES
# ==============================================
class Collaborator(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    matricula = db.Column(db.String(20), unique=True)
    nome = db.Column(db.String(150))
    email = db.Column(db.String(120), unique=True)
    telefone = db.Column(db.String(20))
    setor = db.Column(db.String(120))
    funcao = db.Column(db.String(120))
    foto = db.Column(db.String(255))
    ativo = db.Column(db.Boolean, default=True)
    password = db.Column(db.String(255))
    created_at = db.Column(db.DateTime, default=agora)

class Announcement(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(150), nullable=False)
    message = db.Column(db.Text, nullable=False)
    target_type = db.Column(db.String(20), default="all")  # all | internal | company | user
    target_company_id = db.Column(db.Integer, db.ForeignKey("company.id"), nullable=True)
    target_user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=True)
    sender_id = db.Column(db.Integer, db.ForeignKey("user.id"))
    created_at = db.Column(db.DateTime, default=agora)
    expires_at = db.Column(db.DateTime, nullable=True)  # Null = Permanente

    target_company = db.relationship("Company")
    target_user = db.relationship("User", foreign_keys=[target_user_id])
    sender = db.relationship("User", foreign_keys=[sender_id])

class AnnouncementRead(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    announcement_id = db.Column(db.Integer, db.ForeignKey("announcement.id"), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    read_at = db.Column(db.DateTime, default=agora)

# ==============================================
# MANUAIS
# ==============================================
class Manual(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(150), nullable=False)
    file_path = db.Column(db.String(255), nullable=False)
    target_role = db.Column(db.String(20), default="colaborador")  # admin | colaborador
    created_at = db.Column(db.DateTime, default=agora)

class SystemManual(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    role = db.Column(db.String(20), unique=True)  # admin | colaborador
    content = db.Column(db.Text)
    updated_at = db.Column(db.DateTime, default=agora, onupdate=agora)

# ==============================================
# AUTOMAÇÕES
# ==============================================
class Automation(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(150), nullable=False)
    description = db.Column(db.Text)
    link = db.Column(db.String(255))
    created_at = db.Column(db.DateTime, default=agora)

# ==============================================
# TREINAMENTOS (LMS)
# ==============================================
class Training(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(150), nullable=False)
    description = db.Column(db.Text)
    content_url = db.Column(db.String(255))  # Link vídeo ou PDF (mantido por retrocompatibilidade)
    target_type = db.Column(db.String(20), default="all")  # all | internal | company | user
    target_company_id = db.Column(db.Integer, db.ForeignKey("company.id"), nullable=True)
    target_user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=True)
    created_at = db.Column(db.DateTime, default=agora)
    
    badge_icon = db.Column(db.String(50), default="fa-award")
    badge_color = db.Column(db.String(20), default="#3b82f6")
    allow_retake = db.Column(db.Boolean, default=True)

    target_company = db.relationship("Company")
    target_user = db.relationship("User", foreign_keys=[target_user_id])
    questions = db.relationship("TrainingQuestion", backref="training", cascade="all, delete-orphan")
    modules = db.relationship("TrainingModule", backref="training", cascade="all, delete-orphan", order_by="TrainingModule.order")

class TrainingModule(db.Model):
    __tablename__ = "training_module"
    id = db.Column(db.Integer, primary_key=True)
    training_id = db.Column(db.Integer, db.ForeignKey("training.id"), nullable=False)
    title = db.Column(db.String(150), nullable=False)
    description = db.Column(db.Text)
    video_path = db.Column(db.String(255), nullable=True) # Vídeo enviado (.mp4, etc)
    image_path = db.Column(db.String(255), nullable=True) # Imagem enviada (.png, .jpg, etc)
    video_url = db.Column(db.String(255), nullable=True)  # Link externo (YouTube)
    order = db.Column(db.Integer, default=0)
    created_at = db.Column(db.DateTime, default=agora)

class TrainingQuestion(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    training_id = db.Column(db.Integer, db.ForeignKey("training.id"))
    question_text = db.Column(db.Text, nullable=False)
    options = db.relationship("TrainingOption", backref="question", cascade="all, delete-orphan")

class TrainingOption(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    question_id = db.Column(db.Integer, db.ForeignKey("training_question.id"))
    option_text = db.Column(db.String(255), nullable=False)
    is_correct = db.Column(db.Boolean, default=False)

class TrainingProgress(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"))
    training_id = db.Column(db.Integer, db.ForeignKey("training.id"))
    completed_at = db.Column(db.DateTime, default=agora)

    user = db.relationship("User")
    training = db.relationship("Training")

class TrainingAttempt(db.Model):
    __tablename__ = "training_attempt"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    training_id = db.Column(db.Integer, db.ForeignKey("training.id"), nullable=False)
    score = db.Column(db.Integer, default=0)
    total = db.Column(db.Integer, default=0)
    passed = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=agora)

    user = db.relationship("User")
    training = db.relationship("Training")

# ==============================================
# COLABORADORES EXTERNOS (DOS CLIENTES)
# ==============================================
class ExternalCollaborator(db.Model):
    __tablename__ = "external_collaborator"
    id = db.Column(db.Integer, primary_key=True)
    company_id = db.Column(db.Integer, db.ForeignKey("company.id"))
    name = db.Column(db.String(150), nullable=False)
    email = db.Column(db.String(120), unique=True)
    phone = db.Column(db.String(20))
    created_at = db.Column(db.DateTime, default=agora)

    company = db.relationship("Company")

# =====================================================
# LOGIN MANAGER
# =====================================================
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# =====================================================
# CRIA ADMIN AUTOMÁTICO
# =====================================================
@app.before_request
def criar_admin_automatico():
    admin = User.query.filter_by(email="admin@admin.com").first()
    if not admin:
        novo = User(
            name="Administrador",
            email="admin@admin.com",
            password=generate_password_hash("123"),
            role="admin"
        )
        db.session.add(novo)
        db.session.commit()

# =====================================================
# RESUMO DO DASHBOARD
# =====================================================
def tarefas_por_empresa():
    dados = []
    hoje = date.today().strftime("%Y-%m-%d")

    empresas = Company.query.all()

    for emp in empresas:
        contratos = Contract.query.filter_by(company_id=emp.id, status="ativo").all()
        if not contratos:
            continue

        ids = [c.id for c in contratos]

        total = Task.query.filter(Task.contract_id.in_(ids)).count()
        concluidas = Task.query.filter(
            Task.contract_id.in_(ids),
            Task.status == "concluida"
        ).count()
        pendentes = Task.query.filter(
            Task.contract_id.in_(ids),
            Task.status == "pendente"
        ).count()
        atrasadas = Task.query.filter(
            Task.contract_id.in_(ids),
            Task.status == "pendente",
            Task.due_date < hoje
        ).count()

        # Adições de dados detalhados para o dashboard individual premium
        tarefas_detalhadas = (
            Task.query
            .filter(Task.contract_id.in_(ids))
            .order_by(Task.status.asc(), Task.due_date.asc(), Task.priority.desc())
            .limit(5)
            .all()
        )

        priority_high = Task.query.filter(Task.contract_id.in_(ids), Task.priority.in_(["Alta", "Urgente"])).count()
        priority_normal = Task.query.filter(Task.contract_id.in_(ids), Task.priority == "Normal").count()
        priority_low = Task.query.filter(Task.contract_id.in_(ids), Task.priority == "Baixa").count()

        contratos_detalhes = []
        for c in contratos:
            contratos_detalhes.append({
                "id": c.id,
                "start_date": c.start_date,
                "end_date": c.end_date,
                "description": c.description or "Sem descrição cadastrada."
            })

        dados.append({
            "empresa": emp,
            "contratos": len(contratos),
            "contratos_lista": contratos_detalhes,
            "total": total,
            "concluidas": concluidas,
            "pendentes": pendentes,
            "atrasadas": atrasadas,
            "recentes_tarefas": tarefas_detalhadas,
            "priority_high": priority_high,
            "priority_normal": priority_normal,
            "priority_low": priority_low
        })

    return dados

# =====================================================
# DASHBOARD
# =====================================================
@app.route("/")
@login_required
def dashboard():
    if current_user.role == "colaborador":
        return redirect("/painel-colaborador")
    
    if current_user.role == "cliente_colaborador":
        return redirect("/portal-externo")

    return render_template("dashboard.html", dados=tarefas_por_empresa())

# =====================================================
# LISTA DE CONTRATOS
# =====================================================
@app.route("/contracts")
@login_required
def lista_contratos():
    ordem = request.args.get("ordem", "recent")
    if ordem == "oldest":
        contracts = Contract.query.order_by(Contract.id.asc()).all()
    else:
        contracts = Contract.query.order_by(Contract.id.desc()).all()
    return render_template("contracts.html", contracts=contracts, current_ordem=ordem)

# =====================================================
# NOVO CONTRATO
# =====================================================
@app.route("/contracts/new", methods=["POST"])
@login_required
def new_contract():
    logo_file = request.files.get("logo")

    logo_path = None
    if logo_file and logo_file.filename != "":
        filename = secure_filename(logo_file.filename)

        folder = app.config["LOGO_FOLDER"]
        if not os.path.exists(folder):
            os.makedirs(folder)

        logo_path = os.path.join(folder, filename)
        logo_file.save(logo_path)

    company = Company(
        name=request.form["company"],
        cnpj=request.form["cnpj"],
        logo=logo_path
    )
    db.session.add(company)
    db.session.commit()

    contract = Contract(
        company_id=company.id,
        description=request.form["desc"],
        start_date=request.form["start"],
        end_date=request.form["end"],
        status="ativo"
    )
    db.session.add(contract)
    db.session.commit()

    flash("Contrato criado!", "success")
    return redirect("/contracts")

# =====================================================
# EDITAR CONTRATO (DADOS BÁSICOS / VIGÊNCIA + LOGO)
# =====================================================
@app.route("/contracts/edit/<int:id>", methods=["POST"])
@login_required
def edit_contract(id):
    contrato = Contract.query.get_or_404(id)
    empresa = contrato.company

    # ========= Dados texto =========
    nome = request.form.get("company")
    cnpj = request.form.get("cnpj")
    descricao = request.form.get("desc")

    if nome:
        empresa.name = nome
    if cnpj:
        empresa.cnpj = cnpj
    if descricao is not None:
        contrato.description = descricao

    # ========= Logo (novo) =========
    remover_logo = request.form.get("remover_logo") == "1"
    logo_file = request.files.get("logo")

    # remover logo atual
    if remover_logo:
        try:
            if empresa.logo and os.path.exists(empresa.logo):
                os.remove(empresa.logo)
        except:
            pass
        empresa.logo = None

    # se enviar uma nova logo, substitui
    if logo_file and logo_file.filename != "":
        folder = app.config["LOGO_FOLDER"]
        if not os.path.exists(folder):
            os.makedirs(folder)

        filename = secure_filename(logo_file.filename)
        novo_path = os.path.join(folder, filename)
        logo_file.save(novo_path)

        # apaga logo antiga (se existir)
        try:
            if empresa.logo and os.path.exists(empresa.logo):
                os.remove(empresa.logo)
        except:
            pass

        empresa.logo = novo_path

    db.session.commit()
    flash("Contrato atualizado com sucesso!", "success")
    return redirect("/contracts")


# =====================================================
# ATUALIZAR VIGÊNCIA (EXTENSÃO)
# =====================================================
@app.route("/contracts/vigencia/<int:id>", methods=["POST"])
@login_required
def update_vigencia(id):
    contrato = Contract.query.get_or_404(id)

    novo_inicio = request.form.get("start") or contrato.start_date
    novo_fim = request.form.get("end")

    if not novo_fim:
        flash("Informe a nova data de término da vigência.", "error")
        return redirect("/contracts")

    contrato.start_date = novo_inicio
    contrato.end_date = novo_fim
    db.session.commit()

    flash("Vigência atualizada com sucesso!", "success")
    return redirect("/contracts")

# =====================================================
# REATIVAR CONTRATO ENCERRADO
# =====================================================
@app.route("/contracts/reactivate/<int:id>", methods=["POST"])
@login_required
def reactivate_contract(id):
    contrato = Contract.query.get_or_404(id)

    if contrato.status != "encerrado":
        flash("Somente contratos encerrados podem ser reativados.", "error")
        return redirect("/contracts")

    novo_inicio = request.form.get("start") or contrato.start_date
    novo_fim = request.form.get("end")

    if not novo_fim:
        flash("Informe a nova vigência (pelo menos a data de término).", "error")
        return redirect("/contracts")

    contrato.start_date = novo_inicio
    contrato.end_date = novo_fim
    contrato.status = "ativo"

    db.session.commit()
    flash("✅ Contrato enviado com sucesso!", "success")
    return redirect(f"/contract/{id}")

# =====================================================
# REGISTRAR COLABORADOR EXTERNO (CLIENTE)
# =====================================================
@app.route("/contract/<int:id>/collaborators/add", methods=["POST"])
@login_required
def add_external_collaborator(id):
    admin_only()
    contr = Contract.query.get_or_404(id)
    name = request.form.get("name")
    email = request.form.get("email")
    phone = request.form.get("phone")
    password = request.form.get("password")
    
    # Se não informar senha, usa o CNPJ automático
    if not password:
        raw_cnpj = contr.company.cnpj or "123456"
        password = "".join(filter(str.isdigit, raw_cnpj))
        if not password: password = "123456"

    # Verifica se já existe
    if User.query.filter_by(email=email).first():
        flash("❌ Este e-mail já está em uso!", "error")
        return redirect(f"/contract/{id}")

    # Criar registro de colaborador externo
    novo_ext = ExternalCollaborator(
        company_id=contr.company_id,
        name=name,
        email=email,
        phone=phone
    )
    db.session.add(novo_ext)

    # Criar User correspondente
    usuario = User(
        name=name,
        email=email,
        role="cliente_colaborador",
        company_id=contr.company_id
    )
    usuario.set_password(password)
    db.session.add(usuario)
    db.session.commit()

    flash(f"✅ Colaborador {name} registrado com sucesso!", "success")
    return redirect(f"/contract/{id}")

# =====================================================
# EDITAR COLABORADOR EXTERNO
# =====================================================
@app.route("/contract/<int:id>/collaborators/edit/<int:colab_id>", methods=["POST"])
@login_required
def edit_external_collaborator(id, colab_id):
    admin_only()
    colab = ExternalCollaborator.query.get_or_404(colab_id)
    user = User.query.filter_by(email=colab.email).first()
    
    name = request.form.get("name")
    email = request.form.get("email")
    phone = request.form.get("phone")
    password = request.form.get("password")
    
    colab.name = name
    colab.email = email
    colab.phone = phone
    
    if user:
        user.name = name
        user.email = email
        if password:
            user.set_password(password)
            
    db.session.commit()
    flash(f"✅ Cadastro de {name} atualizado!", "success")
    return redirect(f"/contract/{id}")

@app.route("/contract/<int:id>/collaborators/delete/<int:colab_id>", methods=["POST"])
@login_required
def delete_external_collaborator(id, colab_id):
    admin_only()
    colab = ExternalCollaborator.query.get_or_404(colab_id)
    # Também remove o usuário correspondente se existir
    user = User.query.filter_by(email=colab.email).first()
    if user:
        db.session.delete(user)
    
    db.session.delete(colab)
    db.session.commit()
    return jsonify({"sucesso": True, "mensagem": "Colaborador removido com sucesso."})

# =====================================================
# ENCERRAR CONTRATO
# =====================================================
@app.route("/contracts/end/<int:id>", methods=["POST"])
@login_required
def end_contract(id):
    data = request.get_json()
    senha = data.get("senha")

    if not check_password_hash(current_user.password, senha):
        return jsonify({"sucesso": False, "mensagem": "Senha incorreta"})

    contrato = Contract.query.get_or_404(id)
    contrato.status = "encerrado"
    db.session.commit()

    return jsonify({"sucesso": True})

# =====================================================
# EXCLUIR CONTRATO
# =====================================================
@app.route("/contracts/delete/<int:id>", methods=["POST"])
@login_required
def delete_contract(id):
    contrato = Contract.query.get_or_404(id)

    if contrato.status != "encerrado":
        return jsonify({
            "sucesso": False,
            "mensagem": "Só é possível excluir contratos encerrados"
        })

    arquivos = ContractFile.query.filter_by(contract_id=id).all()
    for arq in arquivos:
        try:
            if arq.file_path and os.path.exists(arq.file_path):
                os.remove(arq.file_path)
        except:
            pass
        db.session.delete(arq)

    db.session.delete(contrato)
    db.session.commit()

    return jsonify({"sucesso": True})

# =====================================================
# VISUALIZAÇÃO DO CONTRATO (COM REORDENAR)
# =====================================================
@app.route("/contract/<int:id>")
@login_required
def contract_view(id):
    contr = Contract.query.get_or_404(id)
    files = ContractFile.query.filter_by(contract_id=id).all()
    
    # Colaboradores da empresa do contrato
    ext_colaboradores = ExternalCollaborator.query.filter_by(company_id=contr.company_id).all()

    colaboradores = User.query.filter_by(role="colaborador").order_by(User.name).all()

    # ✅ pega ordem do select (default: entrega mais próxima)
    ordem = (request.args.get("ordem") or "due_asc").strip()

    # ✅ base query
    if current_user.role == "admin":
        q = Task.query.filter_by(contract_id=id)
    else:
        q = Task.query.filter(
            Task.contract_id == id,
            ((Task.assigned_to == current_user.id) | (Task.assigned_to == None))
        )

    # ✅ REORDENAR
    if ordem == "id_desc":
        q = q.order_by(Task.id.desc())
    elif ordem == "id_asc":
        q = q.order_by(Task.id.asc())
    elif ordem == "due_desc":
        try:
            q = q.order_by(Task.due_date.desc().nullslast(), Task.id.desc())
        except Exception:
            q = q.order_by(Task.due_date.desc(), Task.id.desc())
    else:
        # padrão due_asc (Entrega mais próxima)
        try:
            q = q.order_by(Task.due_date.asc().nullslast(), Task.id.desc())
        except Exception:
            q = q.order_by(Task.due_date.asc(), Task.id.desc())

    tasks = q.all()

    # ⚠️ Se seu Task.due_date NO BANCO estiver como string, o order_by acima
    # pode não ficar perfeito. O ideal é due_date ser coluna Date.
    # Mesmo assim, mantive seu conversor abaixo.
    for t in tasks:
        if isinstance(t.due_date, str) and t.due_date:
            try:
                t.due_date = datetime.strptime(t.due_date, "%Y-%m-%d").date()
            except:
                t.due_date = None

    return render_template(
        "contract_view.html",
        contract=contr,
        files=files,
        tasks=tasks,
        colaboradores=colaboradores,
        ext_colaboradores=ext_colaboradores,
        now=agora,
        ordem_sel=ordem  # ✅ opcional (pra manter marcado, se quiser usar no template)
    )


# =====================================================
# UPLOAD DE ARQUIVOS DO CONTRATO — MÚLTIPLOS (CORRIGIDO)
# Salva em: static/uploads/contratos/
# Grava no banco: static/uploads/contratos/nome.ext
# =====================================================
@app.route("/contract/<int:id>/upload", methods=["POST"])
@login_required
def upload_file(id):
    files = request.files.getlist("files[]")

    if not files:
        flash("Nenhum arquivo enviado!", "error")
        return redirect(f"/contract/{id}")

    # ✅ pasta dentro do static para arquivos de contrato
    contract_subdir = "uploads/contratos"
    contract_upload_folder = os.path.join(app.root_path, "static", contract_subdir)
    os.makedirs(contract_upload_folder, exist_ok=True)

    for file in files:
        if file and file.filename:
            filename = secure_filename(file.filename)

            # salva o arquivo no disco (caminho absoluto)
            save_abs = os.path.join(contract_upload_folder, filename)
            file.save(save_abs)

            # ✅ caminho web para abrir no navegador
            db_path = f"static/{contract_subdir}/{filename}"

            novo = ContractFile(
                contract_id=id,
                file_path=db_path
            )
            db.session.add(novo)

    db.session.commit()
    flash("Arquivo(s) enviado(s) com sucesso!", "success")
    return redirect(f"/contract/{id}")

# =====================================================
# EXCLUIR ARQUIVO DO CONTRATO
# =====================================================
@app.route("/contract/file/delete/<int:file_id>", methods=["POST"])
@login_required
def delete_contract_file(file_id):
    file = ContractFile.query.get_or_404(file_id)
    contract_id = file.contract_id

    try:
        if file.file_path and os.path.exists(file.file_path):
            os.remove(file.file_path)
    except:
        pass

    db.session.delete(file)
    db.session.commit()

    flash("Arquivo removido!", "success")
    return redirect(f"/contract/{contract_id}")

# =====================================================
# TAREFAS – CRIAR
# =====================================================
@app.route("/task/new/<int:contract_id>", methods=["POST"])
@login_required
def new_task(contract_id):
    title = request.form["title"]
    description = request.form.get("description", "")
    due_date = request.form["due_date"]
    priority = request.form.get("priority", "Normal")
    assigned_to_raw = request.form.get("assigned_to")

    assigned_to_value = None if assigned_to_raw in ["all", "", None] else int(assigned_to_raw)

    nova = Task(
        contract_id=contract_id,
        title=title,
        description=description,
        due_date=due_date,
        priority=priority,
        assigned_to=assigned_to_value
    )

    db.session.add(nova)
    db.session.commit()

    flash("Tarefa criada!", "success")
    return redirect(f"/contract/{contract_id}")

# =====================================================
# ✅ NOVAS ROTAS: JSON / EDIT / DELETE
# =====================================================

@app.route("/task/<int:task_id>/json")
@login_required
def task_json(task_id):
    admin_only()

    task = Task.query.get_or_404(task_id)

    # assigned_to: "all" ou id do User
    assigned_to = "all"
    if task.assigned_to is not None:
        assigned_to = str(task.assigned_to)

    # due_date é string "YYYY-MM-DD" no seu banco
    due_date_str = task.due_date or ""

    return jsonify({
        "id": task.id,
        "title": task.title or "",
        "due_date": due_date_str,
        "priority": task.priority or "Normal",
        "description": task.description or "",
        "assigned_to": assigned_to
    })


@app.route("/task/edit/<int:task_id>", methods=["POST"])
@login_required
def task_edit(task_id):
    admin_only()

    task = Task.query.get_or_404(task_id)

    title = request.form.get("title", "").strip()
    due_date = request.form.get("due_date", "").strip()
    priority = request.form.get("priority", "Normal").strip()
    description = request.form.get("description", "").strip()
    assigned_to_raw = request.form.get("assigned_to", "all").strip()

    if not title:
        flash("Título é obrigatório.", "error")
        return redirect(request.referrer or f"/contract/{task.contract_id}")

    # valida formato de data YYYY-MM-DD (mantendo string)
    try:
        datetime.strptime(due_date, "%Y-%m-%d")
    except:
        flash("Data de entrega inválida.", "error")
        return redirect(request.referrer or f"/contract/{task.contract_id}")

    # atualiza campos
    task.title = title
    task.due_date = due_date
    task.priority = priority
    task.description = description

    # responsável
    if assigned_to_raw in ["all", "", None]:
        task.assigned_to = None
    else:
        try:
            uid = int(assigned_to_raw)
        except:
            flash("Responsável inválido.", "error")
            return redirect(request.referrer or f"/contract/{task.contract_id}")

        user = User.query.get(uid)
        if not user or user.role != "colaborador":
            flash("Responsável inválido.", "error")
            return redirect(request.referrer or f"/contract/{task.contract_id}")

        task.assigned_to = user.id

    db.session.commit()
    flash("Tarefa atualizada com sucesso!", "success")
    return redirect(request.referrer or f"/contract/{task.contract_id}")


@app.route("/task/delete/<int:task_id>", methods=["POST"])
@login_required
def task_delete(task_id):
    admin_only()

    task = Task.query.get_or_404(task_id)
    contract_id = task.contract_id

    # 🔥 Remove arquivos físicos vinculados (steps + completion + logs/files)
    # Steps
    for s in TaskStep.query.filter_by(task_id=task.id).all():
        try:
            if s.file_path and os.path.exists(s.file_path):
                os.remove(s.file_path)
        except:
            pass

    # Completion
    completion = TaskCompletion.query.filter_by(task_id=task.id).first()
    if completion:
        try:
            if completion.file_path and os.path.exists(completion.file_path):
                os.remove(completion.file_path)
        except:
            pass

    # Logs + arquivos de logs
    logs = TaskLog.query.filter_by(task_id=task.id).all()
    for lg in logs:
        for f in TaskFile.query.filter_by(task_log_id=lg.id).all():
            try:
                if f.file_path and os.path.exists(f.file_path):
                    os.remove(f.file_path)
            except:
                pass

    # Com cascade no model, deletar Task apaga steps/logs/completion/files do banco
    db.session.delete(task)
    db.session.commit()

    flash("Tarefa excluída com sucesso!", "success")
    return redirect(request.referrer or f"/contract/{contract_id}")

# =====================================================
# TAREFA – REGISTRAR ETAPA (TaskStep)
# =====================================================
@app.route("/task/<int:task_id>/add-step", methods=["POST"])
@login_required
def add_task_step(task_id):
    task = Task.query.get_or_404(task_id)

    description = request.form.get("step_description")
    file = request.files.get("file")

    if not description:
        flash("Descrição da etapa é obrigatória.", "error")
        return redirect(request.referrer or f"/contract/{task.contract_id}")

    file_path = None
    if file and file.filename:
        upload_folder = app.config["UPLOAD_FOLDER"]
        if not os.path.exists(upload_folder):
            os.makedirs(upload_folder)

        filename = secure_filename(file.filename)
        file_path = os.path.join(upload_folder, filename)
        file.save(file_path)

    etapa = TaskStep(
        task_id=task.id,
        user_id=current_user.id,
        description=description,
        file_path=file_path,
        created_at=agora()
    )

    if task.status == "pendente":
        task.status = "andamento"

    db.session.add(etapa)
    db.session.commit()

    flash("Etapa registrada com sucesso!", "success")
    return redirect(request.referrer or f"/contract/{task.contract_id}")

# =====================================================
# LOGS / DETALHES DA TAREFA (CONCLUSÃO + ETAPAS)
# =====================================================
@app.route("/task/logs/<int:task_id>")
@login_required
def task_logs(task_id):
    task = Task.query.get_or_404(task_id)

    completion = TaskCompletion.query.filter_by(task_id=task.id)\
        .order_by(TaskCompletion.created_at.desc()).first()

    conclusao = None
    if completion:
        conclusao = {
            "user": completion.user,
            "data": completion.created_at.strftime("%d/%m/%Y %H:%M"),
            "note": completion.note,
            "file": completion.file_path
        }

    etapas = []
    for e in TaskStep.query.filter_by(task_id=task.id)\
                           .order_by(TaskStep.created_at.asc()).all():
        etapas.append({
            "descricao": e.description,
            "data": e.created_at.strftime("%d/%m/%Y %H:%M"),
            "usuario": e.user.name if e.user else "N/A",
            "arquivo": e.file_path
        })

    return jsonify({
        "status": task.status,
        "conclusao": conclusao,
        "steps": etapas
    })

# =====================================================
# FINALIZAR TAREFA — SUPORTA MÚLTIPLOS ARQUIVOS
# =====================================================
@app.route("/task/complete/<int:task_id>", methods=["POST"])
@login_required
def complete_task(task_id):
    task = Task.query.get_or_404(task_id)
    note = request.form["note"]

    files = request.files.getlist("files[]")

    log = TaskLog(
        task_id=task.id,
        user_id=current_user.id,
        note=note
    )
    db.session.add(log)
    db.session.flush()

    saved_paths = []
    if files:
        if not os.path.exists(app.config["UPLOAD_FOLDER"]):
            os.makedirs(app.config["UPLOAD_FOLDER"])

        for f in files:
            if f and f.filename:
                filename = secure_filename(f.filename)
                save_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
                f.save(save_path)
                saved_paths.append(save_path)

                db.session.add(TaskFile(
                    task_log_id=log.id,
                    file_path=save_path
                ))

    completion = TaskCompletion(
        task_id=task.id,
        user=current_user.name,
        note=note,
        file_path=saved_paths[0] if saved_paths else None
    )
    db.session.add(completion)

    task.status = "concluida"
    task.completed_at = agora()

    db.session.commit()

    flash("Tarefa concluída com sucesso!", "success")

    if current_user.role == "admin":
        return redirect(f"/contract/{task.contract_id}")

    return redirect("/painel-colaborador")

# =====================================================
# SALVAR FOTO DO COLABORADOR
# =====================================================
def salvar_foto(arquivo):
    if not arquivo or arquivo.filename == "":
        return None

    if not os.path.exists(app.config["FOTOS_COLAB"]):
        os.makedirs(app.config["FOTOS_COLAB"])

    filename = secure_filename(arquivo.filename)
    caminho = os.path.join(app.config["FOTOS_COLAB"], filename)
    arquivo.save(caminho)

    return caminho

# =====================================================
# COLABORADORES — LISTA
# =====================================================
@app.route("/colaboradores")
@login_required
def colaboradores():
    if current_user.role != "admin":
        return redirect("/")

    colaboradores = Collaborator.query.order_by(Collaborator.nome).all()
    return render_template("colaboradores.html", colaboradores=colaboradores)

# =====================================================
# COLABORADORES — CRIAR
# =====================================================
@app.route("/colaboradores/salvar", methods=["POST"])
@login_required
def salvar_colaborador():
    if current_user.role != "admin":
        return redirect("/")

    nome = request.form.get("nome")
    matricula = request.form.get("matricula")
    email = request.form.get("email")
    telefone = request.form.get("telefone")
    setor = request.form.get("setor")
    funcao = request.form.get("funcao")
    password = request.form.get("password")

    if Collaborator.query.filter_by(matricula=matricula).first():
        flash("❌ Matrícula já cadastrada!", "error")
        return redirect("/colaboradores")

    if Collaborator.query.filter_by(email=email).first():
        flash("❌ Este e-mail já está em uso (colaborador)!", "error")
        return redirect("/colaboradores")

    if User.query.filter_by(email=email).first():
        flash("❌ Este e-mail já está em uso (usuário)!", "error")
        return redirect("/colaboradores")

    foto = salvar_foto(request.files.get("foto"))

    novo = Collaborator(
        nome=nome,
        matricula=matricula,
        email=email,
        telefone=telefone,
        setor=setor,
        funcao=funcao,
        foto=foto,
        password=generate_password_hash(password)
    )
    db.session.add(novo)
    db.session.commit()

    usuario = User(
        name=nome,
        email=email,
        password=generate_password_hash(password),
        role="colaborador"
    )
    db.session.add(usuario)
    db.session.commit()

    flash("✅ Colaborador cadastrado com sucesso!", "success")
    return redirect("/colaboradores")

# =====================================================
# COLABORADORES — EDITAR
# =====================================================
@app.route("/colaboradores/editar/<int:id>", methods=["POST"])
@login_required
def editar_colaborador(id):
    if current_user.role != "admin":
        return redirect("/")

    colab = Collaborator.query.get_or_404(id)
    email_antigo = colab.email

    colab.nome = request.form.get("nome")
    colab.matricula = request.form.get("matricula")
    colab.email = request.form.get("email")
    colab.telefone = request.form.get("telefone")
    colab.setor = request.form.get("setor")
    colab.funcao = request.form.get("funcao")

    nova_foto = request.files.get("foto")
    if nova_foto and nova_foto.filename != "":
        colab.foto = salvar_foto(nova_foto)

    usuario = User.query.filter_by(email=email_antigo).first()
    if usuario:
        usuario.name = colab.nome
        usuario.email = colab.email

    db.session.commit()

    flash("Atualizado com sucesso!", "success")
    return redirect("/colaboradores")

# =====================================================
# COLABORADORES — EXCLUIR
# =====================================================
@app.route("/colaboradores/excluir/<int:id>", methods=["POST"])
@login_required
def excluir_colaborador(id):
    if current_user.role != "admin":
        return redirect("/")

    colab = Collaborator.query.get_or_404(id)
    usuario = User.query.filter_by(email=colab.email).first()

    if usuario:
        db.session.delete(usuario)

    if colab.foto and os.path.exists(colab.foto):
        try:
            os.remove(colab.foto)
        except:
            pass

    db.session.delete(colab)
    db.session.commit()

    flash("Colaborador removido!", "success")
    return redirect("/colaboradores")

# =====================================================
# ADMINISTRADORES — LISTA
# =====================================================
@app.route("/administradores")
@login_required
def administradores():
    if current_user.role != "admin":
        return redirect("/")
    admins = User.query.filter_by(role="admin").all()
    return render_template("administradores.html", admins=admins)

# =====================================================
# ADMIN — CRIAR
# =====================================================
@app.route("/administradores/salvar", methods=["POST"])
@login_required
def salvar_admin():
    if current_user.role != "admin":
        return redirect("/")

    nome = request.form["nome"]
    email = request.form["email"]
    senha = request.form["senha"]

    if User.query.filter_by(email=email).first():
        flash("❌ Este e-mail já existe!", "error")
        return redirect("/administradores")

    novo_admin = User(
        name=nome,
        email=email,
        password=generate_password_hash(senha),
        role="admin"
    )
    db.session.add(novo_admin)
    db.session.commit()

    flash("Administrador criado!", "success")
    return redirect("/administradores")

# =====================================================
# ADMIN – EDITAR
# =====================================================
@app.route("/administradores/editar/<int:id>", methods=["POST"])
@login_required
def editar_admin(id):
    if current_user.role != "admin":
        return redirect("/")

    admin = User.query.get_or_404(id)

    nome = request.form.get("nome")
    novo_email = request.form.get("email")
    nova_senha = request.form.get("senha")

    if admin.email == "admin@admin.com":
        senha_master = request.form.get("senha_master")

        if not senha_master:
            flash("❌ Senha master obrigatória para editar este administrador.", "error")
            return redirect("/administradores")

        if senha_master != MASTER_PASSWORD:
            flash("❌ Senha master inválida.", "error")
            return redirect("/administradores")

        admin.name = nome

        if nova_senha:
            admin.password = generate_password_hash(nova_senha)

    else:
        if User.query.filter(User.email == novo_email, User.id != id).first():
            flash("❌ Este e-mail já está em uso!", "error")
            return redirect("/administradores")

        admin.name = nome
        admin.email = novo_email

        if nova_senha:
            admin.password = generate_password_hash(nova_senha)

    db.session.commit()

    flash("✅ Administrador atualizado com sucesso!", "success")
    return redirect("/administradores")

# =====================================================
# ADMIN — EXCLUIR
# =====================================================
@app.route("/administradores/excluir/<int:id>", methods=["POST"])
@login_required
def excluir_admin(id):
    if current_user.role != "admin":
        return redirect("/")

    admin = User.query.get_or_404(id)

    if admin.email == "admin@admin.com":
        flash("⚠ Não é permitido excluir o admin principal!", "error")
        return redirect("/administradores")

    db.session.delete(admin)
    db.session.commit()

    flash("Administrador removido!", "success")
    return redirect("/administradores")

# =====================================================
# PAINEL DO COLABORADOR (FILTRO EMPRESA + REORDENAR ENTREGA)
# =====================================================
@app.route("/painel-colaborador")
@login_required
def painel_colaborador():
    if current_user.role != "colaborador":
        return redirect("/")

    # ✅ filtros do GET
    empresa_id = (request.args.get("empresa_id") or "").strip()
    ordem = (request.args.get("ordem") or "due_asc").strip()  # due_asc | due_desc

    # ✅ lista de empresas para o select (somente as que tenham contrato ativo e tarefas do colaborador)
    empresas = (
        Company.query
        .join(Contract, Contract.company_id == Company.id)
        .join(Task, Task.contract_id == Contract.id)
        .filter(Company.active == True)
        .filter(Contract.status == "ativo")
        .filter(Task.status.in_(["pendente", "andamento"]))
        .filter((Task.assigned_to == current_user.id) | (Task.assigned_to == None))
        .distinct()
        .order_by(Company.name.asc())
        .all()
    )

    # ============================
    # QUERY BASE: tarefas do colaborador (pendente/andamento) + contrato ativo + empresa ativa
    # ============================
    query = (
        Task.query
        .join(Contract, Contract.id == Task.contract_id)
        .join(Company, Company.id == Contract.company_id)
        .filter(Company.active == True)
        .filter(Contract.status == "ativo")
        .filter(Task.status.in_(["pendente", "andamento"]))
        .filter((Task.assigned_to == current_user.id) | (Task.assigned_to == None))
    )

    # ============================
    # FILTRO POR EMPRESA
    # ============================
    if empresa_id:
        try:
            query = query.filter(Company.id == int(empresa_id))
        except:
            empresa_id = ""

    # ============================
    # ✅ REORDENAR POR ENTREGA (due_date)
    # ============================
    if ordem == "due_desc":
        # mais distante primeiro
        try:
            query = query.order_by(Task.due_date.desc().nullslast(), Task.id.desc())
        except Exception:
            query = query.order_by(Task.due_date.desc(), Task.id.desc())
    else:
        # padrão: mais próxima primeiro
        try:
            query = query.order_by(Task.due_date.asc().nullslast(), Task.id.desc())
        except Exception:
            query = query.order_by(Task.due_date.asc(), Task.id.desc())

    tarefas = query.all()

    # ============================
    # TRATAR DUE_DATE + ATRASO (serve pro template usar t.is_overdue)
    # ============================
    hoje = date.today()

    for t in tarefas:
        if isinstance(t.due_date, str) and t.due_date:
            try:
                t.due_date = datetime.strptime(t.due_date, "%Y-%m-%d").date()
            except:
                t.due_date = None

        # ✅ atrasada se ainda não concluída e data < hoje
        t.is_overdue = (
            t.status in ["pendente", "andamento"]
            and t.due_date is not None
            and t.due_date < hoje
        )

    # ✅ Conquistas (Selo de Aprovação)
    progressos = TrainingProgress.query.filter_by(user_id=current_user.id).all()
    conquistas = [p.training for p in progressos if p.training]

    return render_template(
        "painel_colaborador.html",
        tarefas=tarefas,
        empresas=empresas,
        conquistas=conquistas,
        empresa_id_sel=str(empresa_id) if empresa_id else "",
        ordem_sel=ordem
    )



@app.route("/alterar_senha", methods=["GET", "POST"])
@login_required
def alterar_senha():
    if request.method == "POST":
        # Suporta tanto JSON quanto Form normal
        if request.is_json:
            data = request.get_json()
            senha_atual = data.get("senha_atual")
            nova_senha = data.get("nova_senha")
            
            if senha_atual and not current_user.check_password(senha_atual):
                return jsonify({"sucesso": False, "mensagem": "Senha atual incorreta."})
            
            current_user.set_password(nova_senha)
            db.session.commit()
            return jsonify({"sucesso": True})
        else:
            nova_senha = request.form.get("nova")
            if not nova_senha:
                flash("❌ A nova senha não pode estar vazia!", "error")
                return redirect("/alterar_senha")
                
            current_user.set_password(nova_senha)
            db.session.commit()
            flash("✅ Senha atualizada com sucesso!", "success")
            
            if current_user.role == 'cliente_colaborador':
                return redirect("/painel_colaborador")
            return redirect("/dashboard")

    return render_template("alterar_senha.html")

# =====================================================
# ATIVIDADES (LOGS) — COM FILTROS + ORDENAR + LOAD MORE
# =====================================================
@app.route("/atividades")
@login_required
def atividades():
    empresa_id = (request.args.get("empresa") or "").strip()
    colaborador_id = (request.args.get("colaborador") or "").strip()
    periodo = (request.args.get("periodo") or "").strip()

    ordem = (request.args.get("ordem") or "recentes").strip()

    # ✅ NOVO: limite só da VISUALIZAÇÃO
    try:
        show = int(request.args.get("show", 20))
    except ValueError:
        show = 20

    # trava de segurança
    if show < 1:
        show = 20
    if show > 500:
        show = 500

    query = (
        TaskLog.query
        .join(Task, Task.id == TaskLog.task_id)
        .join(Contract, Contract.id == Task.contract_id)
        .join(Company, Company.id == Contract.company_id)
        # ✅ não mostrar empresas inativas / contratos inativos
        .filter(Company.active == True)
        .filter(Contract.status == "ativo")
    )
# ✅ join para ordenar por colaborador
    query = query.join(User, User.id == TaskLog.user_id)
    
    # Se não for admin, só vê os próprios logs
    if current_user.role != "admin":
        query = query.filter(TaskLog.user_id == current_user.id)
    
    # filtros
    if empresa_id:
        try:
            query = query.filter(Company.id == int(empresa_id))
        except:
            empresa_id = ""
    
    if colaborador_id:
        try:
            query = query.filter(TaskLog.user_id == int(colaborador_id))
        except:
            colaborador_id = ""
    
    if periodo and " até " in periodo:
        try:
            inicio_str, fim_str = periodo.split(" até ")
    
            data_inicio = datetime.strptime(inicio_str.strip(), "%d/%m/%Y")
            data_fim = datetime.strptime(fim_str.strip(), "%d/%m/%Y")
    
            data_inicio = datetime.combine(data_inicio, time.min)
            data_fim = datetime.combine(data_fim, time.max)
    
            query = query.filter(TaskLog.created_at.between(data_inicio, data_fim))
        except ValueError:
            pass
    
    # ordenação
    if ordem == "antigas":
        query = query.order_by(TaskLog.created_at.asc())
    
    elif ordem == "empresa_az":
        query = query.order_by(Company.name.asc(), TaskLog.created_at.desc())
    
    elif ordem == "empresa_za":
        query = query.order_by(Company.name.desc(), TaskLog.created_at.desc())
    
    elif ordem == "colab_az":
        query = query.order_by(User.name.asc(), TaskLog.created_at.desc())
    
    elif ordem == "colab_za":
        query = query.order_by(User.name.desc(), TaskLog.created_at.desc())
    
    else:  # "recentes"
        query = query.order_by(TaskLog.created_at.desc())
    
    # ✅ AQUI: limita visualização e detecta se tem mais
    rows = query.limit(show + 1).all()
    has_more = len(rows) > show
    logs = rows[:show]
    
    # ✅ Empresas para o select: só ativas e que tenham logs (TaskLog)
    empresas = (
        Company.query
        .join(Contract, Contract.company_id == Company.id)
        .join(Task, Task.contract_id == Contract.id)
        .join(TaskLog, TaskLog.task_id == Task.id)
        .filter(Company.active == True)
        .filter(Contract.status == "ativo")
        .distinct()
        .order_by(Company.name.asc())
        .all()
    )
    
    # ✅ Colaboradores para o select
    # Admin vê todos; colaborador vê só ele mesmo
    if current_user.role == "admin":
        colaboradores = (
            User.query
            .filter_by(role="colaborador")
            .order_by(User.name.asc())
            .all()
        )
    else:
        colaboradores = [current_user]
    
    # ✅ Link "Carregar mais" preservando filtros atuais
    next_url = url_for(
        "atividades",
        empresa=empresa_id or "",
        colaborador=colaborador_id or "",
        periodo=periodo or "",
        ordem=ordem or "recentes",
        show=show + 20,
    )
    
    return render_template(
        "atividades.html",
        logs=logs,
        empresas=empresas,
        colaboradores=colaboradores,
        empresa_sel=empresa_id,
        colaborador_sel=colaborador_id,
        periodo_sel=periodo,
        ordem_sel=ordem,
        show=show,           # ✅ opcional pro template exibir “Exibindo 20 por vez”
        has_more=has_more,   # ✅ pro botão aparecer
        next_url=next_url,   # ✅ link pronto
        total_count=None,
    )

# =====================================================
# LOGIN / LOGOUT
# =====================================================
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form["email"]
        password = request.form["password"]

        user = User.query.filter_by(email=email).first()
        if user and check_password_hash(user.password, password):
            login_user(user)
            return redirect("/")
        else:
            flash("❌ Login inválido!", "error")

    return render_template("login.html")

@app.route("/logout")
def logout():
    logout_user()
    return redirect("/login")

# =====================================================
# RELATÓRIOS (ADMIN) — BASEADO EM CONCLUSÕES (TaskCompletion)
# =====================================================
# =====================================================
# RELATÓRIOS (ADMIN) — BASEADO EM CONCLUSÕES (TaskCompletion)
# =====================================================
@app.route("/relatorios")
@login_required
def relatorios():
    admin_only()

    empresa_id = request.args.get("empresa", "").strip()
    periodo = request.args.get("periodo", "").strip()  # "dd/mm/yyyy até dd/mm/yyyy"
    ordem = (request.args.get("ordem") or "data_desc").strip()

    # ✅ Empresas para o select (somente as que:
    # - estão ativas
    # - possuem contrato ativo
    # - possuem ao menos 1 conclusão (TaskCompletion)
    empresas = (
        Company.query
        .join(Contract, Contract.company_id == Company.id)
        .join(Task, Task.contract_id == Contract.id)
        .join(TaskCompletion, TaskCompletion.task_id == Task.id)
        .filter(Company.active == True)
        .filter(Contract.status == "ativo")
        .distinct()
        .order_by(Company.name.asc())
        .all()
    )

    # ✅ NOVO: limite apenas da VISUALIZAÇÃO
    try:
        show = int(request.args.get("show", 20))
    except ValueError:
        show = 20

    # trava de segurança (ajuste se quiser)
    if show < 1:
        show = 20
    if show > 500:
        show = 500

    # ✅ Empresas para o select (somente as que:
    # - estão ativas
    # - possuem contrato ativo
    # - possuem ao menos 1 conclusão (TaskCompletion)
    empresas = (
        Company.query
        .join(Contract, Contract.company_id == Company.id)
        .join(Task, Task.contract_id == Contract.id)
        .join(TaskCompletion, TaskCompletion.task_id == Task.id)
        .filter(Company.active == True)
        .filter(Contract.status == "ativo")
        .distinct()
        .order_by(Company.name.asc())
        .all()
    )

    # ✅ Base: conclusões (melhor prova de serviço)
    query = (
        TaskCompletion.query
        .join(Task, Task.id == TaskCompletion.task_id)
        .join(Contract, Contract.id == Task.contract_id)
        .join(Company, Company.id == Contract.company_id)
        .filter(Task.status == "concluida")
        .filter(Company.active == True)        # ✅ não traz empresa inativa
        .filter(Contract.status == "ativo")    # ✅ não traz contrato removido/inativo
        .filter(Company.active == True)
        .filter(Contract.status == "ativo")
    )

    # Filtro por empresa
    if empresa_id:
        try:
            query = query.filter(Company.id == int(empresa_id))
        except:
            empresa_id = ""

    # Filtro por período (tz-aware)
    dt_ini, dt_fim = parse_periodo_local(periodo)
    if dt_ini and dt_fim:
        query = query.filter(TaskCompletion.created_at.between(dt_ini, dt_fim))

    # ✅ ORDENAR (conforme o select do template)
    if ordem == "data_asc":
        query = query.order_by(TaskCompletion.created_at.asc())

    elif ordem == "empresa_asc":
        query = query.order_by(Company.name.asc(), TaskCompletion.created_at.desc())

    elif ordem == "empresa_desc":
        query = query.order_by(Company.name.desc(), TaskCompletion.created_at.desc())

    elif ordem == "tarefa_asc":
        query = query.order_by(Task.title.asc(), TaskCompletion.created_at.desc())

    elif ordem == "colab_asc":
        # TaskCompletion.user é string (nome) no seu model
        query = query.order_by(TaskCompletion.user.asc(), TaskCompletion.created_at.desc())

    else:  # padrão: "data_desc"
        query = query.order_by(TaskCompletion.created_at.desc())

    # ✅ AQUI está a mudança: pega 1 a mais para saber se tem mais registros
    rows = query.limit(show + 1).all()
    has_more = len(rows) > show
    logs = rows[:show]

    # ✅ Link "Carregar mais" preservando filtros
    next_show = show + 20
    next_url = url_for(
        "relatorios",
        empresa=str(empresa_id) if empresa_id else "",
        periodo=periodo or "",
        ordem=ordem or "data_desc",
        show=next_show,
    )

    # KPIs (opcional)
    empresas_ativas = Company.query.filter_by(active=True).count()
    contratos_ativos = Contract.query.filter_by(status="ativo").count()
    total_tarefas = Task.query.count()

    return render_template(
        "relatorios.html",
        logs=logs,
        empresas=empresas,
        empresa_sel=str(empresa_id) if empresa_id else "",
        periodo_sel=periodo or "",
        ordem_sel=ordem,
        show=show,                 # ✅ opcional (pra mostrar “Exibindo 20 por vez”)
        has_more=has_more,         # ✅ pro botão aparecer
        next_url=next_url,         # ✅ link pronto pro “Carregar mais”
        total_count=None,          # ✅ se quiser depois eu adiciono count sem pesar
        empresas_ativas=empresas_ativas,
        contratos_ativos=contratos_ativos,
        total_tarefas=total_tarefas
    )



# =====================================================
# EXPORTAR RELATÓRIOS — EXCEL (COMPLETO + FORMATADO)
# =====================================================
@app.route("/relatorios/export/excel")
@login_required
def relatorios_export_excel():
    admin_only()

    import io
    from datetime import datetime

    import pandas as pd
    from flask import send_file
    from openpyxl.utils import get_column_letter
    from openpyxl.styles import Alignment, Font
    from openpyxl.worksheet.table import Table, TableStyleInfo

    empresa_id = (request.args.get("empresa") or "").strip()
    periodo = (request.args.get("periodo") or "").strip()

    query = (
        TaskCompletion.query
        .join(Task, Task.id == TaskCompletion.task_id)
        .join(Contract, Contract.id == Task.contract_id)
        .join(Company, Company.id == Contract.company_id)
        .filter(Task.status == "concluida")
    )

    company_name = "Todas"
    if empresa_id:
        try:
            cid = int(empresa_id)
            query = query.filter(Company.id == cid)
            company = Company.query.get(cid)
            if company:
                company_name = company.name
        except:
            empresa_id = ""
            company_name = "Todas"

    dt_ini, dt_fim = parse_periodo_local(periodo)
    if dt_ini and dt_fim:
        query = query.filter(TaskCompletion.created_at.between(dt_ini, dt_fim))

    items = query.order_by(TaskCompletion.created_at.asc()).all()

    # ---------------------------
    # Monta linhas "completas"
    # ---------------------------
    rows = []
    for c in items:
        task = getattr(c, "task", None)
        contract = getattr(task, "contract", None) if task else None
        company = getattr(contract, "company", None) if contract else None

        created_at = getattr(c, "created_at", None)

        # Campos (padrão)
        empresa_nome = getattr(company, "name", "-") if company else "-"
        contrato_id = getattr(task, "contract_id", None) if task else None

        tarefa_titulo = getattr(task, "title", "-") if task else "-"
        tarefa_desc = getattr(task, "description", None) if task else None
        tarefa_status = getattr(task, "status", None) if task else None

        responsavel = getattr(c, "user", None) or "-"
        observacao = getattr(c, "note", None) or "-"

        # Campos extras (se existirem no seu model)
        contrato_numero = getattr(contract, "number", None) or getattr(contract, "code", None) or "-"
        contrato_inicio = getattr(contract, "start_date", None) or "-"
        contrato_fim = getattr(contract, "end_date", None) or "-"
        contrato_valor = getattr(contract, "value", None) or getattr(contract, "amount", None) or "-"

        prioridade = getattr(task, "priority", None) or "-"
        vencimento = getattr(task, "due_date", None) or getattr(task, "deadline", None) or "-"

        rows.append(
            {
                # Excel entende datetime melhor do que string; então mantenho DateTime real
                "Data/Hora Conclusão": created_at if created_at else None,
                "Empresa": empresa_nome,
                "Contrato ID": contrato_id if contrato_id is not None else "-",
                "Contrato Nº/Código": contrato_numero,
                "Contrato Início": contrato_inicio,
                "Contrato Fim": contrato_fim,
                "Contrato Valor": contrato_valor,
                "Tarefa": tarefa_titulo,
                "Descrição da Tarefa": tarefa_desc if tarefa_desc else "-",
                "Status da Tarefa": tarefa_status if tarefa_status else "-",
                "Prioridade": prioridade,
                "Vencimento": vencimento,
                "Concluído por": responsavel,
                "Observação": observacao,
            }
        )

    df = pd.DataFrame(rows)

    # Aba resumo (se tiver dados)
    resumo_resp = pd.DataFrame()
    resumo_tarefa = pd.DataFrame()
    resumo_mes = pd.DataFrame()

    if not df.empty:
        # Por responsável
        resumo_resp = (
            df.groupby("Concluído por", dropna=False)
            .size()
            .reset_index(name="Qtd")
            .sort_values(["Qtd", "Concluído por"], ascending=[False, True])
        )

        # Por tarefa
        resumo_tarefa = (
            df.groupby("Tarefa", dropna=False)
            .size()
            .reset_index(name="Qtd")
            .sort_values(["Qtd", "Tarefa"], ascending=[False, True])
        )

        # Por mês (usa Data/Hora Conclusão)
        temp = df.copy()
        temp["Mês"] = pd.to_datetime(temp["Data/Hora Conclusão"], errors="coerce").dt.to_period("M").astype(str)
        resumo_mes = (
            temp.groupby("Mês", dropna=False)
            .size()
            .reset_index(name="Qtd")
            .sort_values("Mês", ascending=True)
        )

    # ---------------------------
    # Gera Excel com formatação
    # ---------------------------
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        # Aba principal
        sheet_name = "Relatorio"
        df.to_excel(writer, index=False, sheet_name=sheet_name)

        wb = writer.book
        ws = writer.sheets[sheet_name]

        # Header bonito
        header_font = Font(bold=True)
        for cell in ws[1]:
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

        # Congela cabeçalho
        ws.freeze_panes = "A2"

        # Auto filtro
        ws.auto_filter.ref = ws.dimensions

        # Ajusta largura de colunas + wrap nas colunas longas
        wrap_cols = {"Descrição da Tarefa", "Observação"}
        max_width = 55

        for col_idx, col_name in enumerate(df.columns, start=1):
            letter = get_column_letter(col_idx)

            # define largura
            if col_name in wrap_cols:
                ws.column_dimensions[letter].width = 45
            else:
                # calcula largura baseado nos dados
                values = [str(col_name)] + [str(v) if v is not None else "" for v in df[col_name].head(200).tolist()]
                width = min(max(len(x) for x in values) + 2, max_width)
                ws.column_dimensions[letter].width = max(12, width)

            # alinhamento por coluna
            for row in range(2, ws.max_row + 1):
                cell = ws[f"{letter}{row}"]
                if col_name in wrap_cols:
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
                else:
                    cell.alignment = Alignment(vertical="top", wrap_text=False)

        # Formato de data/hora
        if "Data/Hora Conclusão" in df.columns:
            col_idx = list(df.columns).index("Data/Hora Conclusão") + 1
            col_letter = get_column_letter(col_idx)
            for row in range(2, ws.max_row + 1):
                c = ws[f"{col_letter}{row}"]
                c.number_format = "dd/mm/yyyy hh:mm"

        # Cria tabela do Excel (Table)
        if ws.max_row >= 2 and ws.max_column >= 1:
            tab = Table(displayName="TabelaRelatorio", ref=ws.dimensions)
            style = TableStyleInfo(
                name="TableStyleMedium9",
                showFirstColumn=False,
                showLastColumn=False,
                showRowStripes=True,
                showColumnStripes=False,
            )
            tab.tableStyleInfo = style
            ws.add_table(tab)

        # ---------------------------
        # Aba Resumo
        # ---------------------------
        rs_name = "Resumo"
        ws_r = wb.create_sheet(rs_name)

        ws_r["A1"] = "Resumo do Relatório"
        ws_r["A1"].font = Font(bold=True, size=14)

        ws_r["A3"] = "Empresa:"
        ws_r["B3"] = company_name

        ws_r["A4"] = "Período:"
        ws_r["B4"] = periodo or "Sem filtro"

        ws_r["A5"] = "Emitido em:"
        ws_r["B5"] = datetime.now().strftime("%d/%m/%Y %H:%M")

        ws_r["A6"] = "Total de atividades:"
        ws_r["B6"] = int(len(items))

        # Blocos de tabelas na aba Resumo
        start_row = 8

        def write_df(title, dataframe, start_row):
            ws_r[f"A{start_row}"] = title
            ws_r[f"A{start_row}"].font = Font(bold=True)
            start_row += 1

            if dataframe is None or dataframe.empty:
                ws_r[f"A{start_row}"] = "— Sem dados."
                return start_row + 2

            # header
            for col_idx, col in enumerate(dataframe.columns, start=1):
                cell = ws_r.cell(row=start_row, column=col_idx, value=col)
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal="center")

            # rows
            for r_i, row in enumerate(dataframe.itertuples(index=False), start=start_row + 1):
                for c_i, val in enumerate(row, start=1):
                    ws_r.cell(row=r_i, column=c_i, value=val)

            # ajuste largura
            for col_idx in range(1, len(dataframe.columns) + 1):
                letter = get_column_letter(col_idx)
                ws_r.column_dimensions[letter].width = max(18, min(45, len(str(dataframe.columns[col_idx - 1])) + 10))

            return start_row + len(dataframe) + 3

        start_row = write_df("Atividades por responsável", resumo_resp, start_row)
        start_row = write_df("Atividades por tarefa", resumo_tarefa, start_row)
        start_row = write_df("Atividades por mês", resumo_mes, start_row)

    output.seek(0)
    nome = f"relatorio_{company_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"

    return send_file(
        output,
        as_attachment=True,
        download_name=nome,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )



# =====================================================
# EXPORTAR RELATÓRIOS — WORD (DOCX) NO ESTILO DO MODELO
# =====================================================
@app.route("/relatorios/export/word")
@login_required
def relatorios_export_word():
    admin_only()

    import io
    import os
    from datetime import datetime

    from flask import send_file, current_app
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # ---------------------------
    # ✅ BORDA DA PÁGINA (moldura em volta da folha)
    # ---------------------------
    def add_page_border(section):
        sectPr = section._sectPr

        pgBorders = OxmlElement("w:pgBorders")
        pgBorders.set(qn("w:offsetFrom"), "page")

        for side in ("top", "left", "bottom", "right"):
            element = OxmlElement(f"w:{side}")
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "12")
            element.set(qn("w:space"), "24")
            element.set(qn("w:color"), "000000")
            pgBorders.append(element)

        sectPr.append(pgBorders)

    empresa_id = (request.args.get("empresa") or "").strip()
    periodo = (request.args.get("periodo") or "").strip()

    query = (
        TaskCompletion.query
        .join(Task, Task.id == TaskCompletion.task_id)
        .join(Contract, Contract.id == Task.contract_id)
        .join(Company, Company.id == Contract.company_id)
        .filter(Task.status == "concluida")
    )

    company_name = "Todas"
    if empresa_id:
        try:
            cid = int(empresa_id)
            query = query.filter(Company.id == cid)
            company = Company.query.get(cid)
            if company:
                company_name = company.name
        except:
            empresa_id = ""
            company_name = "Todas"

    dt_ini, dt_fim = parse_periodo_local(periodo)
    if dt_ini and dt_fim:
        query = query.filter(TaskCompletion.created_at.between(dt_ini, dt_fim))

    items = query.order_by(TaskCompletion.created_at.asc()).all()

    # Ano-base (derivado do filtro)
    if dt_ini and dt_fim:
        if dt_ini.year == dt_fim.year:
            ano_base = str(dt_ini.year)
        else:
            ano_base = f"{dt_ini.year}–{dt_fim.year}"
    else:
        ano_base = str(datetime.now().year)

    periodo_str = periodo or "Sem filtro"
    emissao_str = datetime.now().strftime("%d/%m/%Y %H:%M")

    consultoria = "WK Comércio e Consultoria em Segurança do Trabalho LTDA"

    # ---------------------------
    # Helpers DOCX (formatos)
    # ---------------------------
    def set_cell_shading(cell, color_hex="D9D9D9"):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color_hex)
        tcPr.append(shd)

    def set_table_borders(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.append(tblPr)

        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "6")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")
            borders.append(element)
        tblPr.append(borders)

    # ---------------------------
    # Carrega template (se existir)
    # ---------------------------
    template_path = os.path.join(current_app.static_folder, "docs", "RELATORIO_MODELO.docx")
    if os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()

    # ---------------------------
    # ✅ Margens (espaço) + ✅ BORDA DA FOLHA
    # ---------------------------
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)

        add_page_border(section)

    # ---------------------------
    # ✅ Cabeçalho com LOGO (static/logo.png) — menor
    # ---------------------------
    logo_path = os.path.join(current_app.static_folder, "logo.png")
    header = doc.sections[0].header
    for p in header.paragraphs:
        p.text = ""

    if os.path.exists(logo_path):
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(1.8))
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ---------------------------
    # TÍTULO
    # ---------------------------
    title = doc.add_paragraph()
    title_run = title.add_run("RELATÓRIO ANUAL DE ATIVIDADES - SEGURANÇA E SAÚDE DO TRABALHO")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ---------------------------
    # BLOCO INFO
    # ---------------------------
    info = doc.add_table(rows=3, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = True
    set_table_borders(info)

    info.cell(0, 0).text = f"Ano-base: {ano_base}"
    info.cell(0, 1).text = f"Cliente: {company_name}"
    info.cell(1, 0).text = f"Consultoria Responsável: {consultoria}"
    info.cell(1, 1).text = f"Período: {periodo_str}"
    info.cell(2, 0).text = f"Emissão: {emissao_str}"
    info.cell(2, 1).text = f"Total de atividades: {len(items)}"

    for row in info.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()

    # ---------------------------
    # APRESENTAÇÃO (✅ frase EXATA)
    # ---------------------------
    h = doc.add_paragraph()
    h_run = h.add_run("Apresentação")
    h_run.bold = True
    h_run.font.size = Pt(12)

    apresentacao_txt = (
        "Este relatório tem como objetivo apresentar, de forma clara e objetiva, "
        "as atividades realizadas ao longo do ano, conforme o escopo contratual firmado entre as partes, "
        "evidenciando o cumprimento das obrigações legais em Segurança e Saúde do Trabalho (SST), "
        "bem como as ações preventivas implementadas para a melhoria contínua do ambiente laboral"
    )
    p_ap = doc.add_paragraph(apresentacao_txt)
    for r in p_ap.runs:
        r.font.size = Pt(10)

    doc.add_paragraph()

    # ---------------------------
    # TABELA: Atividades Realizadas (✅ removido Responsável)
    # ---------------------------
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("Atividades Realizadas")
    h3_run.bold = True
    h3_run.font.size = Pt(12)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    hdr = table.rows[0].cells
    hdr[0].text = "Data"
    hdr[1].text = "Tarefa"

    for cell in hdr:
        set_cell_shading(cell, "E6E6E6")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)

    for c in items:
        row = table.add_row().cells
        row[0].text = c.created_at.strftime("%d/%m/%Y %H:%M") if c.created_at else "-"
        row[1].text = c.task.title if c.task else "-"

        row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # ✅ Alinhamento: Data centralizada / Tarefa à esquerda (melhor leitura)
        for p in row[0].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(9)

        for p in row[1].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.size = Pt(9)

    buff = io.BytesIO()
    doc.save(buff)
    buff.seek(0)

    nome = f"relatorio_{company_name.replace(' ', '_')}_{ano_base}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    return send_file(
        buff,
        as_attachment=True,
        download_name=nome,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# =====================================================
# EXPORTAR RELATÓRIOS — PDF (ESTILO MODELO ANUAL)
# =====================================================
@app.route("/relatorios/export/pdf")
@login_required
def relatorios_export_pdf():
    admin_only()

    import io
    import os
    from datetime import datetime

    from flask import send_file, current_app
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER  # ✅ necessário para centralizar texto no Paragraph
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image

    empresa_id = (request.args.get("empresa") or "").strip()
    periodo = (request.args.get("periodo") or "").strip()

    query = (
        TaskCompletion.query
        .join(Task, Task.id == TaskCompletion.task_id)
        .join(Contract, Contract.id == Task.contract_id)
        .join(Company, Company.id == Contract.company_id)
        .filter(Task.status == "concluida")
    )

    company_name = "Todas"
    if empresa_id:
        try:
            cid = int(empresa_id)
            query = query.filter(Company.id == cid)
            company = Company.query.get(cid)
            if company:
                company_name = company.name
        except:
            empresa_id = ""
            company_name = "Todas"

    dt_ini, dt_fim = parse_periodo_local(periodo)
    if dt_ini and dt_fim:
        query = query.filter(TaskCompletion.created_at.between(dt_ini, dt_fim))

    items = query.order_by(TaskCompletion.created_at.asc()).all()

    # Ano-base
    if dt_ini and dt_fim:
        if dt_ini.year == dt_fim.year:
            ano_base = str(dt_ini.year)
        else:
            ano_base = f"{dt_ini.year}–{dt_fim.year}"
    else:
        ano_base = str(datetime.now().year)

    total_ativ = len(items)
    periodo_str = periodo or "Sem filtro"
    emissao_str = datetime.now().strftime("%d/%m/%Y %H:%M")
    consultoria = "WK Comércio e Consultoria em Segurança do Trabalho LTDA"

    # ---------------------------
    # ✅ BORDA DA PÁGINA (moldura) no PDF
    # ---------------------------
    def draw_page_border(canvas, doc_):
        canvas.saveState()
        canvas.setStrokeColor(colors.black)
        canvas.setLineWidth(1)
        canvas.rect(
            1.2 * cm,
            1.2 * cm,
            A4[0] - (2.4 * cm),
            A4[1] - (2.4 * cm),
        )
        canvas.restoreState()

    buff = io.BytesIO()
    pdf = SimpleDocTemplate(
        buff,
        pagesize=A4,
        leftMargin=2.2 * cm,
        rightMargin=2.2 * cm,
        topMargin=2.2 * cm,
        bottomMargin=2.2 * cm
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleStyle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        spaceAfter=10,
        alignment=1,
    )
    h2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        spaceBefore=8,
        spaceAfter=6,
    )
    normal = ParagraphStyle(
        "Normal2",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
    )
    small = ParagraphStyle(
        "Small",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        textColor=colors.black,
    )

    # ✅ estilo centralizado para o cabeçalho (Data / Atividade)
    small_center = ParagraphStyle(
        "SmallCenter",
        parent=small,
        alignment=TA_CENTER
    )

    story = []

    # ✅ Logo menor no PDF
    logo_path = os.path.join(current_app.static_folder, "logo.png")
    logo_img = Image(logo_path, width=2.2*cm, height=2.2*cm) if os.path.exists(logo_path) else None

    titulo = Paragraph("RELATÓRIO ANUAL DE ATIVIDADES - SEGURANÇA E SAÚDE DO TRABALHO", title_style)

    if logo_img:
        header = Table([[logo_img, titulo]], colWidths=[2.8*cm, 13.2*cm])
    else:
        header = Table([[titulo]], colWidths=[16.0*cm])

    header.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (0, 0), "LEFT"),
        ("ALIGN", (1, 0), (1, 0), "CENTER"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(header)
    story.append(Spacer(1, 10))

    # Info
    info_data = [
        [Paragraph(f"<b>Ano-base:</b> {ano_base}", normal),
         Paragraph(f"<b>Cliente:</b> {company_name}", normal)],
        [Paragraph(f"<b>Consultoria Responsável:</b> {consultoria}", normal),
         Paragraph(f"<b>Período:</b> {periodo_str}", normal)],
        [Paragraph(f"<b>Emissão:</b> {emissao_str}", normal),
         Paragraph(f"<b>Total de atividades:</b> {total_ativ}", normal)],
    ]
    info = Table(info_data, colWidths=[8.0*cm, 8.0*cm])
    info.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 0.8, colors.black),
        ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(info)
    story.append(Spacer(1, 12))

    # Apresentação (✅ frase EXATA)
    story.append(Paragraph("Apresentação", h2))
    apresentacao_txt = (
        "Este relatório tem como objetivo apresentar, de forma clara e objetiva, as atividades realizadas ao longo do ano, "
        "conforme o escopo contratual firmado entre as partes, evidenciando o cumprimento das obrigações legais em Segurança e "
        "Saúde do Trabalho (SST), bem como as ações preventivas implementadas para a melhoria contínua do ambiente laboral"
    )
    story.append(Paragraph(apresentacao_txt, normal))
    story.append(Spacer(1, 10))

    # Tabela: Atividades Realizadas (✅ removido Responsável)
    story.append(Paragraph("Atividades Realizadas", h2))

    table_rows = [[
        Paragraph("<b>Data</b>", small_center),
        Paragraph("<b>Atividade</b>", small_center),  # ✅ troca "Tarefa" por "Atividade"
    ]]

    for it in items:
        data_str = it.created_at.strftime("%d/%m/%Y %H:%M") if it.created_at else "-"
        atividade = it.task.title if it.task else "-"

        table_rows.append([
            Paragraph(data_str, small),
            Paragraph(atividade, small),
        ])

    # ✅ larguras ajustadas para 2 colunas (alinhado certinho)
    tbl = Table(table_rows, colWidths=[4.2*cm, 11.8*cm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),

        # ✅ garante cabeçalho centralizado no MEIO do campo
        ("VALIGN", (0, 0), (-1, 0), "MIDDLE"),

        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),

        # ✅ um pouco mais de “respiro” para o cabeçalho ficar visualmente no meio
        ("TOPPADDING", (0, 0), (-1, 0), 6),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 6),

        # mantém linhas com padding padrão
        ("TOPPADDING", (0, 1), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 4),
    ]))
    story.append(tbl)

    pdf.build(
        story,
        onFirstPage=draw_page_border,
        onLaterPages=draw_page_border
    )

    buff.seek(0)

    nome = f"relatorio_{ano_base}_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
    return send_file(buff, as_attachment=True, download_name=nome, mimetype="application/pdf")

# =====================================================
# COMUNICADOS, MANUAIS E AUTOMAÇÕES
# =====================================================
@app.route("/comunicados")
@login_required
def lista_comunicados():
    admin_only() # ✅ Apenas admins podem ver a página completa de gestão
    agora_local = agora()
    
    comunicados = Announcement.query.filter(
        (Announcement.expires_at == None) | (Announcement.expires_at > agora_local)
    ).order_by(Announcement.created_at.desc()).all()
    
    manual_admin = SystemManual.query.filter_by(role="admin").first()
    manual_colab = SystemManual.query.filter_by(role="colaborador").first()
    manual_externo = SystemManual.query.filter_by(role="cliente_colaborador").first()
    
    # Fallback se não existirem
    if not manual_admin:
        manual_admin = SystemManual(role="admin", content="")
        db.session.add(manual_admin)
    if not manual_colab:
        manual_colab = SystemManual(role="colaborador", content="")
        db.session.add(manual_colab)
    if not manual_externo:
        manual_externo = SystemManual(role="cliente_colaborador", content="")
        db.session.add(manual_externo)
    
    if db.session.new:
        db.session.commit()
    
    automacoes = Automation.query.order_by(Automation.title.asc()).all()
    # Pega todos os usuários para o admin poder escolher
    all_users = User.query.order_by(User.name).all()
    all_companies = (
        Company.query
        .join(Contract, Contract.company_id == Company.id)
        .filter(Company.active == True, Contract.status == "ativo")
        .distinct()
        .order_by(Company.name)
        .all()
    )
    
    return render_template(
        "comunicados.html",
        comunicados=comunicados,
        manual_admin=manual_admin,
        manual_colab=manual_colab,
        manual_externo=manual_externo,
        automacoes=automacoes,
        all_users=all_users,
        all_companies=all_companies
    )

@app.route("/comunicados/send", methods=["POST"])
@login_required
def send_announcement():
    admin_only()
    
    title = request.form.get("title")
    message = request.form.get("message")
    target_type = request.form.get("target_type")  # all | internal | external | company | user | user_internal | user_external
    target_company_id = request.form.get("target_company_id")
    target_user_id = request.form.get("target_user_id")
    
    # Mapear subtipos para o tipo base 'user'
    if target_type in ["user_internal", "user_external"]:
        target_type = "user"
        
    duration = request.form.get("duration")  # 1 | 3 | 7 | permanent
    
    expires_at = None
    if duration and duration != "permanent":
        expires_at = agora() + timedelta(days=int(duration))
    
    # Validações extras baseadas no tipo
    if target_type == "user" and not target_user_id:
        flash("Selecione um destinatário individual.", "error")
        return redirect("/comunicados")
    if target_type == "company" and not target_company_id:
        flash("Selecione uma empresa específica.", "error")
        return redirect("/comunicados")
        
    novo = Announcement(
        title=title,
        message=message,
        target_type=target_type,
        target_company_id=int(target_company_id) if (target_type == "company" and target_company_id) else None,
        target_user_id=int(target_user_id) if (target_type == "user" and target_user_id) else None,
        sender_id=current_user.id,
        expires_at=expires_at
    )
    db.session.add(novo)
    db.session.commit()
    
    flash("Comunicado enviado com sucesso!", "success")
    return redirect("/comunicados")

@app.route("/manuais/add", methods=["POST"])
@login_required
def add_manual():
    admin_only()
    
    title = request.form.get("title")
    target_role = request.form.get("target_role") or "colaborador"
    file = request.files.get("file")
    
    if not file or file.filename == "":
        flash("Selecione um arquivo para o manual.", "error")
        return redirect("/comunicados")
        
    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config["MANUAIS_FOLDER"], filename)
    file.save(file_path)
    
    novo = Manual(title=title, file_path=file_path, target_role=target_role)
    db.session.add(novo)
    db.session.commit()
    
    flash("Manual adicionado com sucesso!", "success")
    return redirect("/comunicados")

@app.route("/automacoes/add", methods=["POST"])
@login_required
def add_automation():
    admin_only()
    
    title = request.form.get("title")
    description = request.form.get("description")
    link = request.form.get("link")
    
    novo = Automation(title=title, description=description, link=link)
    db.session.add(novo)
    db.session.commit()
    
    flash("Automação adicionada com sucesso!", "success")
    return redirect("/comunicados")

@app.route("/comunicados/delete/<int:id>", methods=["POST"])
@login_required
def delete_announcement(id):
    admin_only()
    item = Announcement.query.get_or_404(id)
    db.session.delete(item)
    db.session.commit()
    return jsonify({"sucesso": True})

@app.route("/manuais/delete/<int:id>", methods=["POST"])
@login_required
def delete_manual(id):
    admin_only()
    item = Manual.query.get_or_404(id)
    try:
        if os.path.exists(item.file_path):
            os.remove(item.file_path)
    except:
        pass
    db.session.delete(item)
    db.session.commit()
    return jsonify({"sucesso": True})

@app.route("/automacoes/delete/<int:id>", methods=["POST"])
@login_required
def delete_automation(id):
    admin_only()
    novo = Automation.query.get_or_404(id)
    db.session.delete(novo)
    db.session.commit()
    return jsonify({"sucesso": True})

# --- APIs PARA TOPBAR (NOTIFICAÇÕES E AJUDA) ---

@app.route("/api/comunicados/recent")
@login_required
def api_comunicados_recent():
    agora_local = agora()
    
    # Lógica de filtro para notificações
    role_filter = Announcement.target_type == "all"
    if current_user.role in ["admin", "colaborador"]:
        role_filter |= (Announcement.target_type == "internal")
    elif current_user.role == "cliente_colaborador":
        role_filter |= (Announcement.target_type == "external")

    # Pega IDs de comunicados já lidos pelo usuário
    lidos = [r.announcement_id for r in AnnouncementRead.query.filter_by(user_id=current_user.id).all()]

    comunicados = Announcement.query.filter(
        (role_filter) | 
        ((Announcement.target_type == "user") & (Announcement.target_user_id == current_user.id)) |
        ((Announcement.target_type == "company") & (Announcement.target_company_id == current_user.company_id)),
        (Announcement.expires_at == None) | (Announcement.expires_at > agora_local),
        ~Announcement.id.in_(lidos) if lidos else True
    ).order_by(Announcement.created_at.desc()).limit(10).all()
    
    lista = []
    for c in comunicados:
        lista.append({
            "id": c.id,
            "title": c.title,
            "message": c.message[:100] + ("..." if len(c.message) > 100 else ""),
            "created_at": c.created_at.strftime("%d/%m/%Y"),
            "sender": c.sender.name if c.sender else "Sistema"
        })
    
    return jsonify({"comunicados": lista})

@app.route("/api/comunicados/read/<int:id>", methods=["POST"])
@login_required
def api_comunicados_read(id):
    # Verifica se já leu
    existente = AnnouncementRead.query.filter_by(announcement_id=id, user_id=current_user.id).first()
    if not existente:
        novo = AnnouncementRead(announcement_id=id, user_id=current_user.id)
        db.session.add(novo)
        db.session.commit()
    return jsonify({"sucesso": True})

@app.route("/api/manuais/save", methods=["POST"])
@login_required
def save_system_manual():
    admin_only()
    data = request.json
    role = data.get("role")
    content = data.get("content")
    
    if role not in ["admin", "colaborador", "cliente_colaborador"]:
        return jsonify({"sucesso": False, "mensagem": "Papel inválido."})
    
    manual = SystemManual.query.filter_by(role=role).first()
    if not manual:
        manual = SystemManual(role=role)
        db.session.add(manual)
    
    manual.content = content
    db.session.commit()
    return jsonify({"sucesso": True})

@app.route("/api/manuais/help")
@login_required
def api_manuais_help():
    if current_user.role == "admin":
        role_buscada = "admin"
    elif current_user.role == "cliente_colaborador":
        role_buscada = "cliente_colaborador"
    else:
        role_buscada = "colaborador"
    manual = SystemManual.query.filter_by(role=role_buscada).first()
    
    if not manual or not manual.content:
        return jsonify({"sucesso": False, "mensagem": "Nenhum manual configurado."})
    
    return jsonify({
        "sucesso": True,
        "content": manual.content
    })

@app.route("/portal-externo")
@login_required
def portal_externo():
    if current_user.role != "cliente_colaborador":
        return redirect("/")
    
    # Cliente Colaborador vê apenas os destinados a todos, os da sua empresa ou os específicos para ele
    treinamentos = Training.query.filter(
        (Training.target_type == "all") |
        ((Training.target_type == "company") & (Training.target_company_id == current_user.company_id)) |
        ((Training.target_type == "user") & (Training.target_user_id == current_user.id))
    ).order_by(Training.created_at.desc()).all()
    
    # Carrega progresso e conquistas
    progressos = TrainingProgress.query.filter_by(user_id=current_user.id).all()
    concluidos_ids = [p.training_id for p in progressos]
    conquistas = [p.training for p in progressos if p.training]
    
    return render_template("portal_externo.html", 
                           treinamentos=treinamentos, 
                           concluidos_ids=concluidos_ids,
                           conquistas=conquistas)

# =====================================================
# TREINAMENTOS (LMS) — LISTA
# =====================================================
@app.route("/treinamentos")
@login_required
def list_treinamentos():
    # Admin vê tudo
    if current_user.role == "admin":
        treinamentos = Training.query.order_by(Training.created_at.desc()).all()
        companies = (
            Company.query
            .join(Contract, Contract.company_id == Company.id)
            .filter(Company.active == True, Contract.status == "ativo")
            .distinct()
            .order_by(Company.name)
            .all()
        )
        users = User.query.filter(User.role != 'admin').order_by(User.name).all()
        concluidos_ids = [p.training_id for p in TrainingProgress.query.all()]
        return render_template("treinamentos.html", treinamentos=treinamentos, companies=companies, users=users, concluidos_ids=concluidos_ids)
    
    # Colaborador Interno vê os destinados a todos, internos ou individuais
    elif current_user.role == "colaborador":
        treinamentos = Training.query.filter(
            (Training.target_type == "all") |
            (Training.target_type == "internal") |
            ((Training.target_type == "user") & (Training.target_user_id == current_user.id))
        ).order_by(Training.created_at.desc()).all()
    
    # Cliente Colaborador vê apenas os destinados a todos, os da sua empresa ou os específicos para ele
    else:
        treinamentos = Training.query.filter(
            (Training.target_type == "all") |
            ((Training.target_type == "company") & (Training.target_company_id == current_user.company_id)) |
            ((Training.target_type == "user") & (Training.target_user_id == current_user.id))
        ).order_by(Training.created_at.desc()).all()
    
    # Pega progresso do usuário
    progressos = TrainingProgress.query.filter_by(user_id=current_user.id).all()
    concluidos_ids = [p.training_id for p in progressos]
    
    # Fallback para variáveis de template
    companies = []
    users = []
    if current_user.role == "admin":
        companies = (
            Company.query
            .join(Contract, Contract.company_id == Company.id)
            .filter(Company.active == True, Contract.status == "ativo")
            .distinct()
            .order_by(Company.name)
            .all()
        )
        users = User.query.filter(User.role != 'admin').order_by(User.name).all()

    return render_template("treinamentos.html", 
                           treinamentos=treinamentos, 
                           concluidos_ids=concluidos_ids,
                           companies=companies,
                           users=users)

# =====================================================
# TREINAMENTOS — SALVAR
# =====================================================
@app.route("/treinamentos/save", methods=["POST"])
@login_required
def save_treinamento():
    admin_only()
    
    id = request.form.get("id")
    title = request.form.get("title")
    description = request.form.get("description")
    badge_icon = request.form.get("badge_icon", "fa-award")
    badge_color = request.form.get("badge_color", "#3b82f6")
    
    allow_retake_raw = request.form.get("allow_retake")
    allow_retake = True
    if allow_retake_raw is not None:
        allow_retake = allow_retake_raw in ["1", "true", "on", True]
    
    if id:
        t = Training.query.get(id)
        t.title = title
        t.description = description
        t.badge_icon = badge_icon
        t.badge_color = badge_color
        t.allow_retake = allow_retake
    else:
        t = Training(
            title=title,
            description=description,
            badge_icon=badge_icon,
            badge_color=badge_color,
            allow_retake=allow_retake,
            target_type="all" # Default is 'all' until distributed
        )
        db.session.add(t)

    db.session.commit()
    
    if request.form.get("ajax") == "1":
        return jsonify({"sucesso": True, "id": t.id})
        
    flash("✅ Treinamento salvo com sucesso!", "success")
    return redirect("/treinamentos")

@app.route("/api/treinamentos/<int:id>/questions")
@login_required
def get_training_questions(id):
    t = Training.query.get_or_404(id)
    questions = []
    for q in t.questions:
        options = []
        for opt in q.options:
            options.append({
                "id": opt.id,
                "text": opt.option_text
            })
        questions.append({
            "id": q.id,
            "text": q.question_text,
            "options": options
        })
    return jsonify(questions)

@app.route("/treinamentos/submit_assessment/<int:id>", methods=["POST"])
@login_required
def submit_assessment(id):
    t = Training.query.get_or_404(id)
    answers = request.json.get("answers") # {question_id: option_id}
    
    # Check if they already attempted and retaking is not allowed
    if not t.allow_retake:
        existing_attempt = TrainingAttempt.query.filter_by(user_id=current_user.id, training_id=t.id).first()
        if existing_attempt:
            return jsonify({
                "sucesso": False, 
                "mensagem": "Você já realizou a prova e este treinamento não permite refazê-la."
            })
            
    total_questions = len(t.questions)
    if total_questions == 0:
        # Se não tem prova, marca como concluído direto
        return complete_training_logic(t.id)

    correct_count = 0
    for q in t.questions:
        user_opt_id = answers.get(str(q.id))
        correct_opt = TrainingOption.query.filter_by(question_id=q.id, is_correct=True).first()
        if correct_opt and str(correct_opt.id) == str(user_opt_id):
            correct_count += 1
            
    passed = (correct_count == total_questions)
    
    # Log the attempt in DB
    attempt = TrainingAttempt(
        user_id=current_user.id,
        training_id=t.id,
        score=correct_count,
        total=total_questions,
        passed=passed
    )
    db.session.add(attempt)
    db.session.commit()
            
    if passed:
        return complete_training_logic(t.id)
    else:
        if not t.allow_retake:
            return jsonify({
                "sucesso": False, 
                "mensagem": f"Você acertou {correct_count} de {total_questions}. Repetição não permitida para esta prova."
            })
        return jsonify({
            "sucesso": False, 
            "mensagem": f"Você acertou {correct_count} de {total_questions}. Tente novamente para ganhar seu selo!"
        })

def complete_training_logic(training_id):
    existing = TrainingProgress.query.filter_by(user_id=current_user.id, training_id=training_id).first()
    if not existing:
        novo = TrainingProgress(user_id=current_user.id, training_id=training_id)
        db.session.add(novo)
        db.session.commit()
    return jsonify({"sucesso": True, "mensagem": "Parabéns! Você concluiu o treinamento e ganhou seu selo!"})

# =====================================================
# TREINAMENTOS — CONCLUIR
# =====================================================
@app.route("/treinamentos/complete/<int:id>", methods=["POST"])
@login_required
def complete_treinamento(id):
    existing = TrainingProgress.query.filter_by(user_id=current_user.id, training_id=id).first()
    if not existing:
        novo = TrainingProgress(user_id=current_user.id, training_id=id)
        db.session.add(novo)
        db.session.commit()
        return jsonify({"sucesso": True, "mensagem": "✅ Parabéns! Treinamento concluído."})
    
    return jsonify({"sucesso": True, "mensagem": "Treinamento já estava concluído."})

# =====================================================
# TREINAMENTOS — EXCLUIR
# =====================================================
@app.route("/treinamentos/delete/<int:id>", methods=["POST"])
@login_required
def delete_treinamento(id):
    admin_only()
    t = Training.query.get_or_404(id)
    
    # Remove progressos associados
    TrainingProgress.query.filter_by(training_id=id).delete()
    
    db.session.delete(t)
    db.session.commit()
    return jsonify({"sucesso": True})

# =====================================================
# TREINAMENTOS — DISTRIBUIR
# =====================================================
@app.route("/treinamentos/distribute/<int:id>", methods=["POST"])
@login_required
def distribute_treinamento(id):
    admin_only()
    t = Training.query.get_or_404(id)
    
    target_type = request.form.get("target_type")
    target_company_id = request.form.get("target_company_id")
    target_user_id = request.form.get("target_user_id")
    
    if target_type == "all":
        t.target_type = "all"
        t.target_company_id = None
        t.target_user_id = None
    elif target_type == "internal":
        t.target_type = "internal"
        t.target_company_id = None
        t.target_user_id = None
    elif target_type == "company":
        t.target_type = "company"
        t.target_company_id = int(target_company_id) if target_company_id and target_company_id != "all" else None
        t.target_user_id = None
    elif target_type == "user":
        t.target_type = "user"
        t.target_company_id = None
        t.target_user_id = int(target_user_id) if target_user_id and target_user_id != "all" else None
        
    db.session.commit()
    flash("✅ Distribuição atualizada com sucesso!", "success")
    return redirect("/treinamentos")

# =====================================================
# TREINAMENTOS — SALVAR QUESTÕES
# =====================================================
@app.route("/treinamentos/questions/save/<int:id>", methods=["POST"])
@login_required
def save_questions(id):
    admin_only()
    t = Training.query.get_or_404(id)
    
    assessment_json = request.form.get("assessment_json")
    if assessment_json:
        import json
        try:
            questions_data = json.loads(assessment_json)
            # Limpa questões antigas e insere novas
            TrainingQuestion.query.filter_by(training_id=t.id).delete()
            for q_data in questions_data:
                if not q_data.get('text'):
                    continue
                q = TrainingQuestion(training_id=t.id, question_text=q_data['text'])
                db.session.add(q)
                db.session.flush() # Para pegar o ID da questão
                for opt_data in q_data['options']:
                    if not opt_data.get('text'):
                        continue
                    opt = TrainingOption(
                        question_id=q.id,
                        option_text=opt_data['text'],
                        is_correct=bool(opt_data.get('is_correct', False))
                    )
                    db.session.add(opt)
            db.session.commit()
            flash("✅ Questionário atualizado com sucesso!", "success")
        except Exception as e:
            db.session.rollback()
            print(f"Erro ao salvar avaliação: {e}")
            flash("❌ Erro ao salvar questionário.", "error")
            
    return redirect("/treinamentos")

# =====================================================
# TREINAMENTOS — API MÓDULOS
# =====================================================
@app.route("/api/treinamentos/<int:id>/modules")
@login_required
def get_training_modules(id):
    t = Training.query.get_or_404(id)
    modules = []
    for m in t.modules:
        modules.append({
            "id": m.id,
            "title": m.title,
            "description": m.description or "",
            "video_path": m.video_path or "",
            "image_path": m.image_path or "",
            "video_url": m.video_url or "",
            "order": m.order
        })
    return jsonify(modules)

# =====================================================
# TREINAMENTOS — SALVAR MÓDULO (WITH UPLOAD)
# =====================================================
@app.route("/treinamentos/module/save", methods=["POST"])
@login_required
def save_module():
    admin_only()
    
    training_id = request.form.get("training_id")
    module_id = request.form.get("module_id")
    title = request.form.get("title")
    description = request.form.get("description")
    video_url = request.form.get("video_url")
    order = request.form.get("order", 0)
    
    if not training_id:
        return jsonify({"sucesso": False, "mensagem": "ID do treinamento é obrigatório."}), 400
        
    t = Training.query.get_or_404(training_id)
    
    # Handle file uploads
    video_file = request.files.get("video_file")
    image_file = request.files.get("image_file")
    
    video_path = None
    image_path = None
    
    folder = app.config["TREINAMENTOS_FOLDER"]
    
    # Save video if uploaded
    if video_file and video_file.filename != "":
        filename = f"vid_{int(datetime.now().timestamp())}_{secure_filename(video_file.filename)}"
        video_path = os.path.join(folder, filename)
        video_file.save(video_path)
        
    # Save image if uploaded
    if image_file and image_file.filename != "":
        filename = f"img_{int(datetime.now().timestamp())}_{secure_filename(image_file.filename)}"
        image_path = os.path.join(folder, filename)
        image_file.save(image_path)

    if module_id:
        m = TrainingModule.query.get(module_id)
        if not m:
            return jsonify({"sucesso": False, "mensagem": "Módulo não encontrado."}), 404
        m.title = title
        m.description = description
        m.video_url = video_url
        m.order = int(order)
        if video_path:
            # Delete old video file if exists
            if m.video_path and os.path.exists(m.video_path):
                try:
                    os.remove(m.video_path)
                except Exception as e:
                    print(f"Erro ao remover vídeo antigo: {e}")
            m.video_path = video_path
        if image_path:
            # Delete old image file if exists
            if m.image_path and os.path.exists(m.image_path):
                try:
                    os.remove(m.image_path)
                except Exception as e:
                    print(f"Erro ao remover imagem antiga: {e}")
            m.image_path = image_path
    else:
        m = TrainingModule(
            training_id=t.id,
            title=title,
            description=description,
            video_url=video_url,
            order=int(order),
            video_path=video_path,
            image_path=image_path
        )
        db.session.add(m)
        
    db.session.commit()
    return jsonify({"sucesso": True, "mensagem": "✅ Módulo salvo com sucesso!"})

# =====================================================
# TREINAMENTOS — EXCLUIR MÓDULO
# =====================================================
@app.route("/treinamentos/module/delete/<int:id>", methods=["POST"])
@login_required
def delete_module(id):
    admin_only()
    m = TrainingModule.query.get_or_404(id)
    
    # Delete physical files
    if m.video_path and os.path.exists(m.video_path):
        try:
            os.remove(m.video_path)
        except Exception as e:
            print(f"Erro ao remover arquivo de vídeo: {e}")
            
    if m.image_path and os.path.exists(m.image_path):
        try:
            os.remove(m.image_path)
        except Exception as e:
            print(f"Erro ao remover arquivo de imagem: {e}")
            
    db.session.delete(m)
    db.session.commit()
    return jsonify({"sucesso": True, "mensagem": "Módulo excluído com sucesso."})

# =====================================================
# EXECUTAR APLICAÇÃO
# =====================================================
if __name__ == "__main__":
    with app.app_context():
        db.create_all()

    # 🔥 Permite acessar de outros PCs na mesma rede
    # Acesse de outro PC: http://IP_DO_SEU_PC:5000
    app.run(host="0.0.0.0", port=5000, debug=True)
